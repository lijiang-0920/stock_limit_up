import os
import sys
import json
import time
import hashlib
import urllib.parse
import requests
from datetime import datetime, timedelta
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
import re
import codecs
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading



# ========== 财联社涨停池相关函数 ==========

def generate_sign(params_dict):
    sorted_data = sorted(params_dict.items(), key=lambda item: item[0])
    query_string = urllib.parse.urlencode(sorted_data)
    sha1_hash = hashlib.sha1(query_string.encode('utf-8')).hexdigest()
    sign = hashlib.md5(sha1_hash.encode('utf-8')).hexdigest()
    return sign

def get_headers():
    return {
        "Host": "x-quote.cls.cn",
        "Connection": "keep-alive",
        "sec-ch-ua": "\"Not A(Brand\";v=\"99\", \"Google Chrome\";v=\"121\", \"Chromium\";v=\"121\"",
        "Accept": "application/json, text/plain, */*",
        "Content-Type": "application/x-www-form-urlencoded",
        "sec-ch-ua-mobile": "?0",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36",
        "sec-ch-ua-platform": "\"Windows\"",
        "Origin": "https://www.cls.cn",
        "Sec-Fetch-Site": "same-site",
        "Sec-Fetch-Mode": "cors",
        "Sec-Fetch-Dest": "empty",
        "Referer": "https://www.cls.cn/",
        "Accept-Encoding": "gzip, deflate, br, zstd",
        "Accept-Language": "zh-CN,zh;q=0.9"
    }

def get_params():
    params = {
        "app": "CailianpressWeb",
        "os": "web",
        "rever": "1",
        "sv": "8.4.6",
        "type": "up_pool",
        "way": "last_px"
    }
    params["sign"] = generate_sign(params)
    return params

def convert_stock_code(code):
    if code.startswith('sh'):
        return code[2:] + '.SH'
    elif code.startswith('sz'):
        return code[2:] + '.SZ'
    return code

def format_plate_names(plates):
    if not plates:
        return ""
    return '|'.join([plate['secu_name'] for plate in plates])

def get_beijing_time():
    """获取北京时间 (UTC+8)"""
    return datetime.utcnow() + timedelta(hours=8)

def fetch_limit_up_data():
    url = "https://x-quote.cls.cn/quote/index/up_down_analysis"
    try:
        params = get_params()
        response = requests.get(url, params=params, headers=get_headers(), timeout=10)
        response.raise_for_status()
        data = response.json()
        if data['code'] != 200:
            print(f"API返回错误: {data['msg']}")
            return None
        return data['data']
    except Exception as e:
        print(f"获取涨停池数据失败: {e}")
        return None

def process_limit_up_data(data):
    if not data:
        return None
    
    current_time = get_beijing_time().strftime("%Y-%m-%d %H:%M:%S")
    current_date = get_beijing_time().strftime("%Y-%m-%d")
    
    formatted_data = []
    for stock in data:
        change_value = float(stock['change']) * 100
        change_percent = f"{change_value:.2f}%"
        
        formatted_stock = {
            "code": convert_stock_code(stock['secu_code']),
            "name": stock['secu_name'].strip(),
            "change_percent": change_percent,
            "price": stock['last_px'],
            "limit_up_time": stock['time'],
            "reason": stock['up_reason'],
            "plates": format_plate_names(stock['plate'])
        }
        formatted_data.append(formatted_stock)
    
    result = {
        "date": current_date,
        "update_time": current_time,
        "count": len(formatted_data),
        "stocks": formatted_data
    }
    return result

def save_limit_up_data(data):
    if not data:
        return
    
    os.makedirs('data', exist_ok=True)
    current_date = data['date']
    
    with open(f'data/{current_date}.json', 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    
    dates = [f.replace('.json', '') for f in os.listdir('data') if f.endswith('.json') and f != 'index.json']
    dates.sort(reverse=True)
    
    with open('data/index.json', 'w', encoding='utf-8') as f:
        json.dump(dates, f, ensure_ascii=False)
    
    print(f"涨停池数据已保存: {current_date}, 共{data['count']}只涨停股")

# ========== 韭研公社文章爬取相关函数 ==========

JIUYAN_USERS = {
    '盘前纪要': {
        'user_url': 'https://www.jiuyangongshe.com/u/4df747be1bf143a998171ef03559b517',
        'user_name': '盘前纪要',
        'save_dir_prefix': '韭研公社_盘前纪要',
        'mode': 'full'
    },
    '盘前解读': {
        'user_url': 'https://www.jiuyangongshe.com/u/97fc2a020e644adb89570e69ae35ec02',
        'user_name': '盘前解读',
        'save_dir_prefix': '韭研公社_盘前解读',
        'mode': 'full'
    },
    '优秀阿呆': {
        'user_url': 'https://www.jiuyangongshe.com/u/88cf268bc56c423c985b87d1b1ff5de4',
        'user_name': '优秀阿呆',
        'save_dir_prefix': '韭研公社_优秀阿呆',
        'mode': 'simple'
    }
}

JIUYAN_HEADERS = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36',
}

def get_target_article_url(user_url, date_str):
    """从用户主页获取指定日期的文章链接"""
    try:
        resp = requests.get(user_url, headers=JIUYAN_HEADERS, timeout=15)
        resp.encoding = resp.apparent_encoding
        soup = BeautifulSoup(resp.text, "html.parser")
        
        for li in soup.find_all('li'):
            title_tag = li.select_one('.book-title span')
            time_tag = li.select_one('.fs13-ash')
            if not title_tag or not time_tag:
                continue
            title = title_tag.text.strip()
            pub_time = time_tag.text.strip()
            if pub_time.startswith(date_str):
                a_tag = li.select_one('a[href^="/a/"]')
                if a_tag:
                    article_url = urljoin(user_url, a_tag['href'])
                    return title, article_url
    except Exception as e:
        print(f"获取文章列表失败: {e}")
    
    return None, None

def fetch_article_content(article_url):
    """获取文章详细内容"""
    try:
        resp = requests.get(article_url, headers=JIUYAN_HEADERS, timeout=15)
        html = resp.text

        pattern = r'content:"(.*?)",url:'
        match = re.search(pattern, html, re.DOTALL)
        if not match:
            return None, None, None

        content_html = match.group(1)
        content_html = content_html.replace('\\\\u', '\\u')
        content_html = codecs.decode(content_html, 'unicode_escape')
        content_html = content_html.encode('latin1').decode('utf-8')
        content_html = content_html.replace('\\"', '"')

        soup = BeautifulSoup(content_html, "html.parser")
        return soup, article_url, html
    except Exception as e:
        print(f"获取文章内容失败: {e}")
        return None, None, None

def save_article_and_generate_json(soup, article_url, save_dir, base_fname, user_info, date_str):
    """保存文章并生成JSON数据"""
    mode = user_info.get('mode', 'full')
    
    # 创建目录
    os.makedirs(save_dir, exist_ok=True)
    
    # 处理图片
    images_data = []
    if mode == 'full':
        img_folder = os.path.join(save_dir, "images")
        os.makedirs(img_folder, exist_ok=True)
        headers_with_referer = JIUYAN_HEADERS.copy()
        headers_with_referer['Referer'] = article_url

        img_counter = 1
        processed_images = set()  # 用于跟踪已处理的图片URL，避免重复
        
        for img in soup.find_all('img'):
            src = img.get('src')
            if not src:
                continue
            if not src.startswith('http'):
                src = urljoin(article_url, src)
            
            # 跳过已经处理过的图片
            if src in processed_images:
                # 如果图片已经处理过，直接替换为已有的占位符
                existing_placeholder = None
                for img_data in images_data:
                    if img_data.get('original_src') == src:
                        existing_placeholder = img_data['placeholder']
                        break
                if existing_placeholder:
                    img.replace_with(existing_placeholder)
                continue
            
            try:
                r = requests.get(src, headers=headers_with_referer, timeout=10)
                if r.status_code != 200:
                    continue
                
                ext = os.path.splitext(urlparse(src).path)[-1]
                if not ext or len(ext) > 5:
                    ext = '.jpg'
                fname = f'img{img_counter}{ext}'
                img_path = os.path.join(img_folder, fname)
                
                with open(img_path, 'wb') as f:
                    f.write(r.content)
                
                # 验证图片
                try:
                    from PIL import Image
                    with Image.open(img_path) as im:
                        im.verify()
                except:
                    if os.path.exists(img_path):
                        os.remove(img_path)
                    continue
                
                # 记录图片信息
                placeholder = f"[图片:img{img_counter}{ext}]"
                images_data.append({
                    "placeholder": placeholder,
                    "filename": fname,
                    "src": f"articles/{user_info['save_dir_prefix']}/{date_str}/images/{fname}",
                    "alt": f"图片{img_counter}",
                    "caption": "",
                    "original_src": src  # 记录原始URL用于去重
                })
                
                # 替换当前img标签为占位符
                img.replace_with(placeholder)
                processed_images.add(src)
                img_counter += 1
                
            except Exception as e:
                print(f"下载图片失败: {e}")
                continue

    # 提取文本内容
    if mode == 'full':
        lines = []
        for block in soup.find_all(['p', 'div', 'li']):
            text = ''.join(block.stripped_strings)
            if text.strip():
                lines.append(text)
        content_text = '\n'.join(lines)
    else:
        lines = []
        for p in soup.find_all('p'):
            text = ''.join(p.stripped_strings)
            if text.strip():
                lines.append(text)
        content_text = '\n\n'.join(lines)

    # 去除重复的图片占位符
    def remove_duplicate_image_placeholders(text):
        import re
        pattern = r'\[图片:([^\]]+)\]'
        seen_images = set()
        
        def replace_func(match):
            img_name = match.group(1)
            if img_name in seen_images:
                return ''  # 移除重复的占位符
            else:
                seen_images.add(img_name)
                return match.group(0)  # 保留第一次出现的占位符
        
        # 先移除重复的占位符
        deduplicated = re.sub(pattern, replace_func, text)
        
        # 清理多余的空行
        lines = deduplicated.split('\n')
        cleaned_lines = []
        prev_empty = False
        
        for line in lines:
            is_empty = not line.strip()
            if is_empty and prev_empty:
                continue  # 跳过连续的空行
            cleaned_lines.append(line)
            prev_empty = is_empty
        
        return '\n'.join(cleaned_lines)

    content_text = remove_duplicate_image_placeholders(content_text)

    # 保存文本文件
    txt_path = os.path.join(save_dir, f"{base_fname}.txt")
    with open(txt_path, 'w', encoding='utf-8-sig') as f:
        f.write(content_text)

    # 保存Word文档（如果是full模式）
    docx_path = None
    if mode == 'full':
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # 处理内容，将图片占位符替换为实际图片
            content_for_docx = content_text
            img_placeholder_pattern = r'(\[图片:[^\]]+\])'
            
            # 按段落和图片占位符分割内容
            parts = re.split(img_placeholder_pattern, content_for_docx)
            
            for part in parts:
                part = part.strip()
                if not part:
                    continue
                
                # 检查是否是图片占位符
                img_match = re.fullmatch(r'\[图片:(.*?)\]', part)
                if img_match:
                    img_filename = img_match.group(1)
                    img_path = os.path.join(img_folder, img_filename)
                    if os.path.exists(img_path):
                        try:
                            # 检查图片尺寸，避免过大
                            from PIL import Image
                            with Image.open(img_path) as pil_img:
                                width, height = pil_img.size
                                # 限制最大宽度为6英寸
                                max_width = Inches(6)
                                if width > height:
                                    doc.add_picture(img_path, width=max_width)
                                else:
                                    # 对于高图，限制高度
                                    max_height = Inches(8)
                                    doc.add_picture(img_path, height=max_height)
                        except Exception as img_error:
                            print(f"插入图片到Word文档失败 {img_filename}: {img_error}")
                            doc.add_paragraph(f'[图片插入失败: {img_filename}]')
                    else:
                        doc.add_paragraph(f'[图片文件不存在: {img_filename}]')
                else:
                    # 普通文本段落
                    if part:
                        doc.add_paragraph(part)
            
            docx_path = os.path.join(save_dir, f"{base_fname}.docx")
            doc.save(docx_path)
            
        except ImportError:
            print("警告：未安装python-docx，跳过Word文档生成")
        except Exception as e:
            print(f"生成Word文档失败: {e}")

    # 清理images_data，移除用于去重的辅助字段
    for img_data in images_data:
        if 'original_src' in img_data:
            del img_data['original_src']

    return {
        "content": content_text,
        "images": images_data,
        "files": {
            "txt": f"articles/{user_info['save_dir_prefix']}/{date_str}/{base_fname}.txt",
            "docx": f"articles/{user_info['save_dir_prefix']}/{date_str}/{base_fname}.docx" if docx_path else None
        },
        "word_count": len(content_text.replace('\n', '').replace(' ', '')),
        "image_count": len(images_data)
    }

def crawl_jiuyan_article(user_key, date_str=None):
    """爬取单个用户的文章 - 移除重试逻辑，由GitHub Actions控制"""
    if user_key not in JIUYAN_USERS:
        print(f"未找到用户配置: {user_key}")
        return None
    
    user_info = JIUYAN_USERS[user_key]
    
    if date_str:
        try:
            target_date = datetime.strptime(date_str, '%Y-%m-%d')
        except:
            print("日期格式错误")
            return None
    else:
        # 使用北京时区的当前日期
        target_date = get_beijing_time()
    
    date_str = target_date.strftime('%Y-%m-%d')
    
    try:
        title, article_url = get_target_article_url(user_info['user_url'], date_str)
        
        if not title:
            print(f"未找到{user_info['user_name']} {date_str}的文章")
            return None

        print(f"找到文章：{title}")
        
        safe_title = re.sub(r'[\\/:*?"<>|]', '_', title)
        
        # 确定保存路径
        if user_key == '优秀阿呆':
            save_dir = os.path.join('articles', user_info['save_dir_prefix'])
        else:
            save_dir = os.path.join('articles', user_info['save_dir_prefix'], date_str)

        soup, article_url, _ = fetch_article_content(article_url)
        if soup is None:
            print("获取文章内容失败")
            return None

        # 保存文章并获取数据
        article_data = save_article_and_generate_json(soup, article_url, save_dir, safe_title, user_info, date_str)
        
        # 构建完整的文章信息
        result = {
            "author": user_info['user_name'],
            "title": title,
            "publish_time": get_beijing_time().strftime("%H:%M"),
            "date": date_str,
            "url": article_url,
            **article_data
        }
        
        print(f"成功保存 {user_info['user_name']} 的文章: {title}")
        return result
        
    except Exception as e:
        print(f"爬取 {user_info['user_name']} 文章时发生错误: {e}")
        return None

def save_articles_index(articles_data, date_str):
    """保存文章索引数据"""
    os.makedirs('articles', exist_ok=True)
    
    # 读取现有索引
    index_file = 'articles/index.json'
    if os.path.exists(index_file):
        with open(index_file, 'r', encoding='utf-8') as f:
            try:
                index_data = json.load(f)
                # 如果是旧格式，转换为新格式
                if 'users' in index_data or 'recent_articles' in index_data:
                    # 提取日期数据
                    new_index_data = {}
                    for key, value in index_data.items():
                        if key.startswith('20') and isinstance(value, dict) and 'articles' in value:
                            new_index_data[key] = value
                    index_data = new_index_data
            except:
                index_data = {}
    else:
        index_data = {}
    
    # 更新当日数据
    if date_str not in index_data:
        index_data[date_str] = {
            "date": date_str,
            "update_time": get_beijing_time().strftime("%Y-%m-%d %H:%M:%S"),
            "articles": []
        }
    
    # 更新或添加文章数据
    existing_articles = index_data[date_str].get("articles", [])
    
    for new_article in articles_data:
        # 检查是否已存在相同作者的文章
        updated = False
        for i, existing_article in enumerate(existing_articles):
            if existing_article.get("author") == new_article.get("author"):
                existing_articles[i] = new_article
                updated = True
                break
        
        # 如果不存在，添加新文章
        if not updated:
            existing_articles.append(new_article)
    
    index_data[date_str]["articles"] = existing_articles
    index_data[date_str]["update_time"] = get_beijing_time().strftime("%Y-%m-%d %H:%M:%S")
    
    # 保存索引
    with open(index_file, 'w', encoding='utf-8') as f:
        json.dump(index_data, f, ensure_ascii=False, indent=2)
    
    print(f"文章索引已更新: {date_str}")

def crawl_single_jiuyan_user(user_key, date_str=None):
    """爬取单个用户的文章"""
    print(f"开始爬取 {user_key} 的文章...")
    
    if user_key not in JIUYAN_USERS:
        print(f"错误：未找到用户 '{user_key}'")
        print(f"可用用户: {', '.join(JIUYAN_USERS.keys())}")
        return None
    
    article_data = crawl_jiuyan_article(user_key, date_str)
    
    if article_data:
        current_date = date_str or get_beijing_time().strftime('%Y-%m-%d')
        save_articles_index([article_data], current_date)
        print(f"成功爬取 {user_key} 的文章")
        return article_data
    else:
        print(f"未能获取 {user_key} 的文章")
        return None

def crawl_all_jiuyan_articles(date_str=None):
    """爬取所有韭研公社文章"""
    print("开始爬取所有韭研公社文章...")
    articles_data = []
    
    for user_key in JIUYAN_USERS.keys():
        print(f"\n处理: {user_key}")
        article_data = crawl_jiuyan_article(user_key, date_str)
        if article_data:
            articles_data.append(article_data)
        time.sleep(2)  # 避免请求过于频繁
    
    # 保存文章索引
    if articles_data:
        current_date = date_str or get_beijing_time().strftime('%Y-%m-%d')
        save_articles_index(articles_data, current_date)
    
    print(f"\n韭研公社文章爬取完成！成功: {len(articles_data)}/{len(JIUYAN_USERS)}")
    return articles_data    

# ========== 韭研公社异动解析相关函数 ==========

def fetch_stock_analysis_data(date_str=None):
    """获取韭研公社异动解析数据"""
    if not date_str:
        date_str = get_beijing_time().strftime('%Y-%m-%d')
    
    url = "https://app.jiuyangongshe.com/jystock-app/api/v1/action/field"
    
    headers = {
        "Host": "app.jiuyangongshe.com",
        "Connection": "keep-alive",
        "sec-ch-ua": '"Not A(Brand";v="99", "Google Chrome";v="121", "Chromium";v="121"',
        "sec-ch-ua-mobile": "?0",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36",
        "Content-Type": "application/json",
        "Accept": "application/json, text/plain, */*",
        "timestamp": str(int(time.time() * 1000)),
        "platform": "3",
        "token": "c9f25f21829e88387639723f4f98272a",
        "sec-ch-ua-platform": '"Windows"',
        "Origin": "https://www.jiuyangongshe.com",
        "Sec-Fetch-Site": "same-site",
        "Sec-Fetch-Mode": "cors",
        "Sec-Fetch-Dest": "empty",
        "Referer": "https://www.jiuyangongshe.com/",
        "Accept-Encoding": "gzip, deflate, br, zstd",
        "Accept-Language": "zh-CN,zh;q=0.9",
        "Cookie": "SESSION=Njk3YTRhNTUtZjFkMi00NzNiLTk2NGYtNTVlNDU5NTRmNWU3; Hm_lvt_58aa18061df7855800f2a1b32d6da7f4=1754989369,1755050145; Hm_lpvt_58aa18061df7855800f2a1b32d6da7f4=1755052203"
    }
    
    payload = {
        "date": date_str,
        "pc": 1
    }
    
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=30)
        response.raise_for_status()
        data = response.json()
        
        if data.get("errCode") == "0" and "data" in data:
            return data["data"]
        else:
            print(f"异动解析API返回错误: {data.get('errCode', '未知错误')}")
            return None
            
    except Exception as e:
        print(f"获取异动解析数据失败: {e}")
        return None

def process_stock_analysis_data(raw_data, date_str):
    """处理异动解析数据"""
    if not raw_data:
        return None
    
    current_time = get_beijing_time().strftime("%Y-%m-%d %H:%M:%S")
    
    categories = []
    total_stocks = 0
    
    for category in raw_data:
        category_name = category.get("name", "")
        category_reason = category.get("reason", "")
        stock_list = category.get("list", [])
        
        if stock_list:  # 只处理有股票数据的分类
            processed_stocks = []
            
            for stock in stock_list:
                stock_code = stock.get("code", "")
                stock_name = stock.get("name", "")
                
                # 获取异动信息
                action_info = stock.get("article", {}).get("action_info", {})
                limit_time = action_info.get("time", "")
                analysis = action_info.get("expound", "")
                
                processed_stock = {
                    "code": stock_code,
                    "name": stock_name,
                    "limit_time": limit_time,
                    "analysis": analysis
                }
                processed_stocks.append(processed_stock)
                total_stocks += 1
            
            category_data = {
                "name": category_name,
                "reason": category_reason,
                "stock_count": len(processed_stocks),
                "stocks": processed_stocks
            }
            categories.append(category_data)
    
    result = {
        "date": date_str,
        "update_time": current_time,
        "category_count": len(categories),
        "total_stocks": total_stocks,
        "categories": categories
    }
    
    return result

def generate_analysis_text_content(data):
    """生成异动解析的文本内容"""
    content = f"韭研公社异动解析 - {data['date']}\n"
    content += f"更新时间: {data['update_time']}\n"
    content += f"板块数量: {data['category_count']} 个\n"
    content += f"股票数量: {data['total_stocks']} 只\n"
    content += "=" * 80 + "\n\n"
    
    for category in data['categories']:
        content += f"=== {category['name']} ===\n"
        if category['reason']:
            content += f"板块异动解析: {category['reason']}\n"
        content += f"涉及股票: {category['stock_count']} 只\n\n"
        
        for stock in category['stocks']:
            content += f"{stock['name']}（{stock['code']}）\n"
            if stock['limit_time']:
                content += f"涨停时间: {stock['limit_time']}\n"
            content += f"个股异动解析: {stock['analysis']}\n"
            content += "\n" + "-" * 80 + "\n\n"
    
    return content

def save_stock_analysis_data(data):
    """保存异动解析数据"""
    if not data:
        return
    
    # 创建数据目录
    os.makedirs('analysis', exist_ok=True)
    current_date = data['date']
    
    # 保存JSON数据
    json_path = f'analysis/{current_date}.json'
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    
    # 保存文本格式
    text_content = generate_analysis_text_content(data)
    txt_path = f'analysis/{current_date}.txt'
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(text_content)
    
    # 更新索引文件
    index_path = 'analysis/index.json'
    if os.path.exists(index_path):
        with open(index_path, 'r', encoding='utf-8') as f:
            try:
                index_data = json.load(f)
                if not isinstance(index_data, dict):
                    index_data = {}
            except:
                index_data = {}
    else:
        index_data = {}
    
    # 更新索引
    index_data[current_date] = {
        "date": current_date,
        "update_time": data['update_time'],
        "category_count": data['category_count'],
        "total_stocks": data['total_stocks'],
        "files": {
            "json": f"analysis/{current_date}.json",
            "txt": f"analysis/{current_date}.txt"
        }
    }
    
    with open(index_path, 'w', encoding='utf-8') as f:
        json.dump(index_data, f, ensure_ascii=False, indent=2)
    
    print(f"异动解析数据已保存: {current_date}, 共{data['category_count']}个板块，{data['total_stocks']}只股票")

def crawl_stock_analysis(date_str=None):
    """爬取异动解析数据"""
    print("开始获取韭研公社异动解析数据...")
    
    if not date_str:
        date_str = get_beijing_time().strftime('%Y-%m-%d')
    
    try:
        raw_data = fetch_stock_analysis_data(date_str)
        processed_data = process_stock_analysis_data(raw_data, date_str)
        
        if processed_data:
            save_stock_analysis_data(processed_data)
            print(f"异动解析数据获取成功: {date_str}")
            return processed_data
        else:
            print(f"未获取到 {date_str} 的异动解析数据")
            return None
            
    except Exception as e:
        print(f"获取异动解析数据时发生错误: {e}")
        return None

# ========== 通达信龙虎榜相关函数 ==========

def get_tdx_lhb_overview():
    """获取通达信龙虎榜总览数据"""
    url = "https://fk.tdx.com.cn/TQLEX?Entry=CWServ.tdxsj_lhbd_lhbzl"
    
    headers = {
        "Host": "fk.tdx.com.cn",
        "Connection": "keep-alive",
        "Content-Type": "application/x-www-form-urlencoded; charset=UTF-8",
        "X-Requested-With": "XMLHttpRequest",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36",
        "Origin": "https://fk.tdx.com.cn",
        "Referer": "https://fk.tdx.com.cn/site/tdxsj/html/tdxsj_lhbd.html?from=www&webfrom=1&pc=0",
        "Accept-Encoding": "gzip, deflate, br, zstd",
        "Accept-Language": "zh-CN,zh;q=0.9",
        "Accept": "text/plain, */*; q=0.01",
        "sec-ch-ua": "\"Not A(Brand\";v=\"99\", \"Google Chrome\";v=\"121\", \"Chromium\";v=\"121\"",
        "sec-ch-ua-mobile": "?0",
        "sec-ch-ua-platform": "\"Windows\""
    }
    
    data = {"Params": ["0", "0", "0"]}
    
    try:
        response = requests.post(url, headers=headers, data=json.dumps(data), timeout=15)
        response.raise_for_status()
        
        result = response.json()
        if result.get('ErrorCode') == 0:
            return result
        else:
            print(f"龙虎榜API返回错误: {result.get('ErrorCode')}")
            return None
    except Exception as e:
        print(f"获取龙虎榜总览失败: {e}")
        return None

def parse_lhb_overview(overview_data):
    """解析龙虎榜总览数据"""
    if not overview_data or not overview_data.get('ResultSets'):
        return None
    
    result_sets = overview_data['ResultSets']
    stocks_info = []
    trading_date = None
    
    # 解析股票列表
    if len(result_sets) > 0 and result_sets[0].get('Count', 0) > 0:
        table0 = result_sets[0]
        for row in table0['Content']:
            stock_info = {
                "code": row[0],
                "name": row[1],
                "market_code": row[2],
                "change_percent": float(row[3]) if row[3] else 0,
                "close_price": float(row[4]) if row[4] else 0,
                "market_name": row[5]
            }
            stocks_info.append(stock_info)
    
    # 解析交易日期
    if len(result_sets) > 1 and result_sets[1].get('Count', 0) > 0:
        table1 = result_sets[1]
        if table1['Content']:
            trading_date = table1['Content'][0][2]
    
    return {
        "trading_date": trading_date,
        "total_count": len(stocks_info),
        "stocks": stocks_info
    }

def get_single_dragon_tiger_detail(stock_code, date):
    """获取单只股票的龙虎榜详细信息"""
    url = "https://fk.tdx.com.cn/TQLEX?Entry=CWServ.tdxsj_lhbd_ggxq"
    
    headers = {
        'Host': 'fk.tdx.com.cn',
        'Connection': 'keep-alive',
        'sec-ch-ua': '"Not A(Brand";v="99", "Google Chrome";v="121", "Chromium";v="121"',
        'Accept': 'text/plain, */*; q=0.01',
        'Content-Type': 'application/x-www-form-urlencoded; charset=UTF-8',
        'X-Requested-With': 'XMLHttpRequest',
        'sec-ch-ua-mobile': '?0',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36',
        'sec-ch-ua-platform': '"Windows"',
        'Origin': 'https://fk.tdx.com.cn',
        'Sec-Fetch-Site': 'same-origin',
        'Sec-Fetch-Mode': 'cors',
        'Sec-Fetch-Dest': 'empty',
        'Referer': f'https://fk.tdx.com.cn/site/tdxsj/html/tdxsj_lhbd_ggxq.html?back=tdxsj_lhbd,%E9%BE%99%E8%99%8E%E6%A6%9C%E4%B8%AA%E8%82%A1,{stock_code}&pc=0&webfrom=1',
        'Accept-Encoding': 'gzip, deflate, br, zstd',
        'Accept-Language': 'zh-CN,zh;q=0.9'
    }
    
    payload = json.dumps({"Params": ["1", stock_code, date]})
    
    try:
        response = requests.post(url, headers=headers, data=payload, timeout=10)
        response.raise_for_status()
        
        raw_data = response.json()
        
        if raw_data.get('ErrorCode') != 0:
            return {"code": stock_code, "status": "api_error", "error_code": raw_data.get('ErrorCode')}
        
        result_sets = raw_data.get('ResultSets', [])
        
        structured_data = {
            "code": stock_code,
            "query_date": date,
            "status": "success"
        }
        
        # 解析基本信息
        if len(result_sets) > 0 and result_sets[0].get('Count', 0) > 0:
            table0 = result_sets[0]
            basic_info = table0['Content'][0]
            
            structured_data["lhb_info"] = {
                "list_code": basic_info[0],
                "list_reason": basic_info[1],
                "volume": float(basic_info[2]) if basic_info[2] else 0,
                "amount": float(basic_info[3]) if basic_info[3] else 0,
                "close_price": float(basic_info[4]) if basic_info[4] else 0,
                "change_percent": float(basic_info[5]) if basic_info[5] else 0
            }
        else:
            structured_data["status"] = "no_detail_data"
            return structured_data
        
        # 解析席位信息
        if len(result_sets) > 1 and result_sets[1].get('Count', 0) > 0:
            table1 = result_sets[1]
            
            buy_seats = []
            sell_seats = []
            
            for row in table1['Content']:
                seat_info = {
                    "rank": int(row[0]) if row[0] else 0,
                    "department_name": row[2],
                    "buy_amount": float(row[3]) if row[3] else 0,
                    "sell_amount": float(row[4]) if row[4] else 0,
                    "net_amount": float(row[5]) if row[5] else 0,
                    "direction": row[7],
                    "label": row[12] if row[12] else ""
                }
                
                # 计算占比
                total_amount = basic_info[3] if basic_info[3] > 0 else 1
                seat_info["amount_ratio"] = round(abs(row[5]) / total_amount * 100, 2)
                
                if row[7] == "B":
                    buy_seats.append(seat_info)
                else:
                    sell_seats.append(seat_info)
            
            structured_data["buy_seats"] = buy_seats
            structured_data["sell_seats"] = sell_seats
            
            # 计算资金流向
            buy_total = sum([seat.get("buy_amount", 0) for seat in buy_seats])
            sell_total = sum([seat.get("sell_amount", 0) for seat in sell_seats])
            
            structured_data["capital_flow"] = {
                "buy_total": buy_total,
                "sell_total": sell_total,
                "net_inflow": buy_total - sell_total,
                "buy_ratio": round(buy_total / total_amount * 100, 2),
                "sell_ratio": round(sell_total / total_amount * 100, 2)
            }
        
        return structured_data
        
    except Exception as e:
        return {"code": stock_code, "status": "query_failed", "error": str(e)}

def crawl_dragon_tiger_data(date_str=None, max_workers=5, delay=0.1):
    """爬取龙虎榜数据"""
    print("开始获取通达信龙虎榜数据...")
    
    if not date_str:
        date_str = get_beijing_time().strftime('%Y-%m-%d')
    
    try:
        # 获取总览数据
        print("1. 获取龙虎榜总览...")
        overview_data = get_tdx_lhb_overview()
        if not overview_data:
            print("获取龙虎榜总览失败")
            return None
        
        # 解析总览数据
        print("2. 解析总览数据...")
        parsed_overview = parse_lhb_overview(overview_data)
        if not parsed_overview:
            print("解析龙虎榜总览失败")
            return None
        
        trading_date = parsed_overview["trading_date"]
        stocks_list = parsed_overview["stocks"]
        
        print(f"发现 {len(stocks_list)} 只龙虎榜股票，交易日期: {trading_date}")
        
        # 并发获取详细数据
        print(f"3. 并发获取详细数据（线程数: {max_workers}）...")
        
        all_detailed_data = {
            "date": trading_date,
            "update_time": get_beijing_time().strftime("%Y-%m-%d %H:%M:%S"),
            "total_count": len(stocks_list),
            "overview": parsed_overview,
            "details": {},
            "statistics": {
                "success_count": 0,
                "failed_count": 0,
                "no_detail_count": 0
            }
        }
        
        completed_count = 0
        lock = threading.Lock()
        
        def query_with_delay(stock_info):
            nonlocal completed_count
            
            time.sleep(delay)
            detail_data = get_single_dragon_tiger_detail(stock_info["code"], trading_date)
            
            # 合并总览信息
            detail_data.update({
                "name": stock_info["name"],
                "market_name": stock_info["market_name"],
                "overview_change_percent": stock_info["change_percent"],
                "overview_close_price": stock_info["close_price"]
            })
            
            with lock:
                completed_count += 1
                all_detailed_data["details"][stock_info["code"]] = detail_data
                
                status = detail_data.get("status", "unknown")
                if status == "success":
                    all_detailed_data["statistics"]["success_count"] += 1
                    lhb_info = detail_data.get("lhb_info", {})
                    flow_info = detail_data.get("capital_flow", {})
                    print(f"[{completed_count:2d}/{len(stocks_list)}] ✓ {stock_info['code']} {stock_info['name'][:8]} - "
                          f"{lhb_info.get('list_reason', 'N/A')[:15]} - 净流入: {flow_info.get('net_inflow', 0):>8.2f}万元")
                elif status == "no_detail_data":
                    all_detailed_data["statistics"]["no_detail_count"] += 1
                    print(f"[{completed_count:2d}/{len(stocks_list)}] - {stock_info['code']} {stock_info['name'][:8]} - 无详细数据")
                else:
                    all_detailed_data["statistics"]["failed_count"] += 1
                    print(f"[{completed_count:2d}/{len(stocks_list)}] ✗ {stock_info['code']} {stock_info['name'][:8]} - 查询失败")
            
            return detail_data
        
        # 使用线程池并发执行
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_stock = {
                executor.submit(query_with_delay, stock): stock 
                for stock in stocks_list
            }
            
            for future in as_completed(future_to_stock):
                try:
                    future.result()
                except Exception as e:
                    stock = future_to_stock[future]
                    print(f"✗ {stock['code']} 查询异常: {e}")
        
        # 保存数据
        save_dragon_tiger_data(all_detailed_data)
        
        success_rate = all_detailed_data["statistics"]["success_count"] / len(stocks_list) * 100
        print(f"龙虎榜数据获取完成: {trading_date}, 成功率: {success_rate:.1f}%")
        
        return all_detailed_data
        
    except Exception as e:
        print(f"获取龙虎榜数据时发生错误: {e}")
        return None

def save_dragon_tiger_data(data):
    """保存龙虎榜数据"""
    if not data:
        return
    
    # 创建数据目录
    os.makedirs('dragon_tiger', exist_ok=True)
    current_date = data['date']
    
    # 保存JSON数据
    json_path = f'dragon_tiger/{current_date}.json'
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    
    # 生成文本格式
    text_content = generate_dragon_tiger_text_content(data)
    txt_path = f'dragon_tiger/{current_date}.txt'
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(text_content)
    
    # 更新索引文件
    index_path = 'dragon_tiger/index.json'
    if os.path.exists(index_path):
        with open(index_path, 'r', encoding='utf-8') as f:
            try:
                index_data = json.load(f)
                if not isinstance(index_data, dict):
                    index_data = {}
            except:
                index_data = {}
    else:
        index_data = {}
    
    # 更新索引
    index_data[current_date] = {
        "date": current_date,
        "update_time": data['update_time'],
        "total_count": data['total_count'],
        "success_count": data['statistics']['success_count'],
        "success_rate": round(data['statistics']['success_count'] / data['total_count'] * 100, 1),
        "files": {
            "json": f"dragon_tiger/{current_date}.json",
            "txt": f"dragon_tiger/{current_date}.txt"
        }
    }
    
    with open(index_path, 'w', encoding='utf-8') as f:
        json.dump(index_data, f, ensure_ascii=False, indent=2)
    
    print(f"龙虎榜数据已保存: {current_date}, 共{data['total_count']}只股票，成功{data['statistics']['success_count']}只")

def generate_dragon_tiger_text_content(data):
    """生成龙虎榜的文本内容"""
    content = f"通达信龙虎榜数据 - {data['date']}\n"
    content += f"更新时间: {data['update_time']}\n"
    content += f"股票总数: {data['total_count']} 只\n"
    content += f"查询成功: {data['statistics']['success_count']} 只\n"
    content += f"成功率: {data['statistics']['success_count']/data['total_count']*100:.1f}%\n"
    content += "=" * 80 + "\n\n"
    
    # 市场分布统计
    market_stats = {}
    for stock in data['overview']['stocks']:
        market = stock['market_name']
        if market not in market_stats:
            market_stats[market] = 0
        market_stats[market] += 1
    
    content += "=== 市场分布 ===\n"
    for market, count in market_stats.items():
        content += f"{market}: {count}只\n"
    content += "\n"
    
    # 详细数据
    content += "=== 龙虎榜详细数据 ===\n\n"
    
    for stock_code, detail in data['details'].items():
        if detail.get('status') != 'success':
            continue
            
        content += f"股票代码: {stock_code}\n"
        content += f"股票名称: {detail.get('name', 'N/A')}\n"
        content += f"市场: {detail.get('market_name', 'N/A')}\n"
        
        lhb_info = detail.get('lhb_info', {})
        content += f"收盘价: {lhb_info.get('close_price', 0):.2f}元\n"
        content += f"涨跌幅: {lhb_info.get('change_percent', 0):.2f}%\n"
        content += f"上榜原因: {lhb_info.get('list_reason', 'N/A')}\n"
        content += f"成交额: {lhb_info.get('amount', 0):.2f}万元\n"
        content += f"成交量: {lhb_info.get('volume', 0):.2f}万股\n"
        
        # 资金流向
        flow_info = detail.get('capital_flow', {})
        if flow_info:
            content += f"买入合计: {flow_info.get('buy_total', 0):.2f}万元\n"
            content += f"卖出合计: {flow_info.get('sell_total', 0):.2f}万元\n"
            content += f"净流入: {flow_info.get('net_inflow', 0):.2f}万元\n"
        
        # 买入席位
        buy_seats = detail.get('buy_seats', [])
        if buy_seats:
            content += "\n买入席位TOP5:\n"
            for i, seat in enumerate(buy_seats[:5], 1):
                content += f"  {i}. {seat.get('department_name', 'N/A')} - "
                content += f"买入: {seat.get('buy_amount', 0):.2f}万元 "
                content += f"占比: {seat.get('amount_ratio', 0):.2f}% "
                content += f"{seat.get('label', '')}\n"
        
        # 卖出席位
        sell_seats = detail.get('sell_seats', [])
        if sell_seats:
            content += "\n卖出席位TOP5:\n"
            for i, seat in enumerate(sell_seats[:5], 1):
                content += f"  {i}. {seat.get('department_name', 'N/A')} - "
                content += f"卖出: {seat.get('sell_amount', 0):.2f}万元 "
                content += f"占比: {seat.get('amount_ratio', 0):.2f}% "
                content += f"{seat.get('label', '')}\n"
        
        content += "\n" + "-" * 80 + "\n\n"
    
    return content


# ========== 主函数和统一接口 ==========

def main_limit_up():
    """主函数 - 财联社涨停池数据"""
    try:
        print("开始获取财联社涨停池数据...")
        raw_data = fetch_limit_up_data()
        processed_data = process_limit_up_data(raw_data)
        
        if processed_data:
            save_limit_up_data(processed_data)
            print("财联社涨停池数据处理完成！")
        else:
            print("没有获取到有效数据")
    except Exception as e:
        print(f"处理涨停池数据时发生错误: {e}")

def main():
    """主函数 - 根据命令行参数决定执行哪个功能"""
    if len(sys.argv) == 1:
        # 默认执行涨停池数据获取
        main_limit_up()
        generate_all_pages()
    elif len(sys.argv) >= 2:
        command = sys.argv[1].lower()
        
        if command == 'limitup':
            main_limit_up()
            
        elif command == 'jiuyan':
            if len(sys.argv) == 2:
                crawl_all_jiuyan_articles()
            elif len(sys.argv) == 3:
                user_key = sys.argv[2]
                crawl_single_jiuyan_user(user_key)
            elif len(sys.argv) == 4:
                user_key = sys.argv[2]
                date_str = sys.argv[3]
                crawl_single_jiuyan_user(user_key, date_str)
        
        elif command == 'analysis':
            if len(sys.argv) == 2:
                crawl_stock_analysis()
            elif len(sys.argv) == 3:
                date_str = sys.argv[2]
                crawl_stock_analysis(date_str)
        
        elif command == 'dragon_tiger':  # 添加龙虎榜命令
            if len(sys.argv) == 2:
                crawl_dragon_tiger_data()
            elif len(sys.argv) == 3:
                date_str = sys.argv[2]
                crawl_dragon_tiger_data(date_str)
                
        elif command == 'all':
            print("执行所有功能...")
            main_limit_up()
            print("\n" + "="*60 + "\n")
            crawl_all_jiuyan_articles()
            print("\n" + "="*60 + "\n")
            crawl_stock_analysis()
            print("\n" + "="*60 + "\n")
            crawl_dragon_tiger_data()  # 添加龙虎榜
            
        else:
            print("使用说明:")
            print("  python script.py                           # 默认获取涨停池数据并生成网页")
            print("  python script.py limitup                   # 获取涨停池数据")
            print("  python script.py jiuyan                    # 爬取韭研公社所有用户文章")
            print("  python script.py jiuyan 盘前纪要           # 爬取韭研公社指定用户文章")
            print("  python script.py jiuyan 盘前纪要 2025-01-21 # 爬取韭研公社指定用户指定日期文章")
            print("  python script.py analysis                  # 获取异动解析数据")
            print("  python script.py analysis 2025-01-21       # 获取指定日期异动解析数据")
            print("  python script.py dragon_tiger              # 获取龙虎榜数据")
            print("  python script.py dragon_tiger 2025-01-21   # 获取指定日期龙虎榜数据")
            print("  python script.py all                       # 执行所有功能")
            print("\n可用的韭研公社用户:")
            for key, info in JIUYAN_USERS.items():
                print(f"  {key} - {info['user_name']}")

if __name__ == "__main__":
    main()























