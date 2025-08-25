import json
import requests
import urllib.parse
import hashlib
import os
import re
import codecs
import sys
import time
from datetime import datetime, timedelta
from urllib.parse import urljoin, urlparse
from bs4 import BeautifulSoup

# ========== 财联社涨停池相关函数 ==========

def generate_sign(params_dict):
    sorted_data = sorted(params_dict.items(), key=lambda item: item[0])
    query_string = urllib.parse.urlencode(sorted_data)
    sha1_hash = hashlib.sha1(query_string.encode('utf-8')).hexdigest()
    sign = hashlib.md5(sha1_hash.encode('utf-8')).hexdigest()
    return sign

def get_headers():
    return {
        "Host": "x-quote.cls.cn",
        "Connection": "keep-alive",
        "sec-ch-ua": "\"Not A(Brand\";v=\"99\", \"Google Chrome\";v=\"121\", \"Chromium\";v=\"121\"",
        "Accept": "application/json, text/plain, */*",
        "Content-Type": "application/x-www-form-urlencoded",
        "sec-ch-ua-mobile": "?0",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36",
        "sec-ch-ua-platform": "\"Windows\"",
        "Origin": "https://www.cls.cn",
        "Sec-Fetch-Site": "same-site",
        "Sec-Fetch-Mode": "cors",
        "Sec-Fetch-Dest": "empty",
        "Referer": "https://www.cls.cn/",
        "Accept-Encoding": "gzip, deflate, br, zstd",
        "Accept-Language": "zh-CN,zh;q=0.9"
    }

def get_params():
    params = {
        "app": "CailianpressWeb",
        "os": "web",
        "rever": "1",
        "sv": "8.4.6",
        "type": "up_pool",
        "way": "last_px"
    }
    params["sign"] = generate_sign(params)
    return params

def convert_stock_code(code):
    if code.startswith('sh'):
        return code[2:] + '.SH'
    elif code.startswith('sz'):
        return code[2:] + '.SZ'
    return code

def format_plate_names(plates):
    if not plates:
        return ""
    return '|'.join([plate['secu_name'] for plate in plates])

def get_beijing_time():
    return datetime.now() 

def fetch_limit_up_data():
    url = "https://x-quote.cls.cn/quote/index/up_down_analysis"
    try:
        params = get_params()
        response = requests.get(url, params=params, headers=get_headers(), timeout=10)
        response.raise_for_status()
        data = response.json()
        if data['code'] != 200:
            print(f"API返回错误: {data['msg']}")
            return None
        return data['data']
    except Exception as e:
        print(f"获取涨停池数据失败: {e}")
        return None

def process_limit_up_data(data):
    if not data:
        return None
    
    current_time = get_beijing_time().strftime("%Y-%m-%d %H:%M:%S")
    current_date = get_beijing_time().strftime("%Y-%m-%d")
    
    formatted_data = []
    for stock in data:
        change_value = float(stock['change']) * 100
        change_percent = f"{change_value:.2f}%"
        
        formatted_stock = {
            "code": convert_stock_code(stock['secu_code']),
            "name": stock['secu_name'].strip(),
            "change_percent": change_percent,
            "price": stock['last_px'],
            "limit_up_time": stock['time'],
            "reason": stock['up_reason'],
            "plates": format_plate_names(stock['plate'])
        }
        formatted_data.append(formatted_stock)
    
    result = {
        "date": current_date,
        "update_time": current_time,
        "count": len(formatted_data),
        "stocks": formatted_data
    }
    return result

def save_limit_up_data(data):
    if not data:
        return
    
    os.makedirs('data', exist_ok=True)
    current_date = data['date']
    
    with open(f'data/{current_date}.json', 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    
    dates = [f.replace('.json', '') for f in os.listdir('data') if f.endswith('.json') and f != 'index.json']
    dates.sort(reverse=True)
    
    with open('data/index.json', 'w', encoding='utf-8') as f:
        json.dump(dates, f, ensure_ascii=False)
    
    print(f"涨停池数据已保存: {current_date}, 共{data['count']}只涨停股")

# ========== 韭研公社文章爬取相关函数（修复版）==========

JIUYAN_USERS = {
    '盘前纪要': {
        'user_url': 'https://www.jiuyangongshe.com/u/4df747be1bf143a998171ef03559b517',
        'user_name': '盘前纪要',
        'save_dir_prefix': '韭研公社_盘前纪要',
        'default_time': '07:40',
        'retry_time': '08:00',
        'mode': 'full'
    },
    '盘前解读': {
        'user_url': 'https://www.jiuyangongshe.com/u/97fc2a020e644adb89570e69ae35ec02',
        'user_name': '盘前解读',
        'save_dir_prefix': '韭研公社_盘前解读',
        'default_time': '08:17',
        'retry_time': '08:27',
        'mode': 'full'
    },
    '优秀阿呆': {
        'user_url': 'https://www.jiuyangongshe.com/u/88cf268bc56c423c985b87d1b1ff5de4',
        'user_name': '优秀阿呆',
        'save_dir_prefix': '韭研公社_优秀阿呆',
        'default_time': '23:30',
        'retry_time': None,
        'mode': 'simple'
    }
}

# 增强的请求头
JIUYAN_HEADERS = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36',
    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
    'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
    'Accept-Encoding': 'gzip, deflate, br',
    'Connection': 'keep-alive',
    'Upgrade-Insecure-Requests': '1',
    'Cache-Control': 'max-age=0'
}

def get_target_article_url(user_url, date_str):
    """获取目标文章URL - 修复版，增加重试机制和错误处理"""
    session = requests.Session()
    session.headers.update(JIUYAN_HEADERS)
    
    max_retries = 3
    for attempt in range(max_retries):
        try:
            print(f"尝试第 {attempt + 1} 次请求韭研公社用户页面...")
            
            # 增加超时时间并启用重定向
            resp = session.get(user_url, timeout=30, allow_redirects=True)
            
            # 检查响应状态
            if resp.status_code != 200:
                print(f"请求失败，状态码: {resp.status_code}")
                if attempt < max_retries - 1:
                    time.sleep(2 * (attempt + 1))  # 递增等待时间
                    continue
                else:
                    return None, None
            
            # 检查内容长度
            if len(resp.content) < 1000:
                print(f"响应内容过短，可能是错误页面，长度: {len(resp.content)}")
                if attempt < max_retries - 1:
                    time.sleep(2 * (attempt + 1))
                    continue
                else:
                    return None, None
            
            print(f"成功获取页面，内容长度: {len(resp.content)}")
            
            # 设置正确的编码
            resp.encoding = resp.apparent_encoding or 'utf-8'
            
            # 解析页面
            soup = BeautifulSoup(resp.text, "html.parser")
            
            # 查找文章列表
            article_found = False
            for li in soup.find_all('li'):
                title_tag = li.select_one('.book-title span')
                time_tag = li.select_one('.fs13-ash')
                
                if not title_tag or not time_tag:
                    continue
                
                title = title_tag.text.strip()
                pub_time = time_tag.text.strip()
                
                # 调试输出
                if attempt == 0:  # 只在第一次尝试时输出
                    print(f"找到文章: {title} - 发布时间: {pub_time}")
                
                if pub_time.startswith(date_str):
                    a_tag = li.select_one('a[href^="/a/"]')
                    if a_tag:
                        article_url = urljoin(user_url, a_tag['href'])
                        print(f"找到目标文章: {title}")
                        return title, article_url
                    article_found = True
            
            if not article_found and attempt == 0:
                print(f"未找到 {date_str} 的文章")
            
            # 如果没找到文章但请求成功，不需要重试
            return None, None
            
        except requests.exceptions.Timeout:
            print(f"请求超时（第 {attempt + 1} 次）")
            if attempt < max_retries - 1:
                time.sleep(3 * (attempt + 1))
                continue
                
        except requests.exceptions.ConnectionError as e:
            print(f"连接错误（第 {attempt + 1} 次）: {e}")
            if attempt < max_retries - 1:
                time.sleep(5 * (attempt + 1))
                continue
                
        except Exception as e:
            print(f"获取文章列表时发生错误（第 {attempt + 1} 次）: {e}")
            if attempt < max_retries - 1:
                time.sleep(2 * (attempt + 1))
                continue
    
    print("所有重试都失败了")
    return None, None

def fetch_article_content(article_url):
    """获取文章内容 - 增强版"""
    session = requests.Session()
    session.headers.update(JIUYAN_HEADERS)
    
    try:
        print(f"获取文章内容: {article_url}")
        resp = session.get(article_url, timeout=30)
        
        if resp.status_code != 200:
            print(f"获取文章内容失败，状态码: {resp.status_code}")
            return None, None, None
            
        html = resp.text

        pattern = r'content:"(.*?)",url:'
        match = re.search(pattern, html, re.DOTALL)
        if not match:
            print("未找到文章内容")
            return None, None, None

        content_html = match.group(1)
        content_html = content_html.replace('\\\\u', '\\u')
        content_html = codecs.decode(content_html, 'unicode_escape')
        content_html = content_html.encode('latin1').decode('utf-8')
        content_html = content_html.replace('\\"', '"')

        soup = BeautifulSoup(content_html, "html.parser")
        return soup, article_url, html
        
    except Exception as e:
        print(f"获取文章内容时发生错误: {e}")
        return None, None, None

def save_article_and_generate_json(soup, article_url, save_dir, base_fname, user_info, date_str):
    """保存文章并生成JSON数据"""
    mode = user_info.get('mode', 'full')
    
    # 创建目录
    os.makedirs(save_dir, exist_ok=True)
    
    # 处理图片
    images_data = []
    if mode == 'full':
        img_folder = os.path.join(save_dir, "images")
        os.makedirs(img_folder, exist_ok=True)
        headers_with_referer = JIUYAN_HEADERS.copy()
        headers_with_referer['Referer'] = article_url

        img_counter = 1
        for img in soup.find_all('img'):
            src = img.get('src')
            if not src:
                continue
            if not src.startswith('http'):
                src = urljoin(article_url, src)
            
            try:
                r = requests.get(src, headers=headers_with_referer, timeout=10)
                if r.status_code != 200:
                    continue
                
                ext = os.path.splitext(urlparse(src).path)[-1]
                if not ext or len(ext) > 5:
                    ext = '.jpg'
                fname = f'img{img_counter}{ext}'
                img_path = os.path.join(img_folder, fname)
                
                with open(img_path, 'wb') as f:
                    f.write(r.content)
                
                # 验证图片
                try:
                    from PIL import Image
                    with Image.open(img_path) as im:
                        im.verify()
                except:
                    if os.path.exists(img_path):
                        os.remove(img_path)
                    continue
                
                # 记录图片信息
                placeholder = f"[图片:img{img_counter}{ext}]"
                images_data.append({
                    "placeholder": placeholder,
                    "filename": fname,
                    "src": f"articles/{user_info['save_dir_prefix']}/{date_str}/images/{fname}",
                    "alt": f"图片{img_counter}",
                    "caption": ""
                })
                
                img.replace_with(placeholder)
                img_counter += 1
                
            except Exception as e:
                print(f"下载图片失败: {e}")
                continue

    # 提取文本内容
    if mode == 'full':
        lines = []
        for block in soup.find_all(['p', 'div', 'li']):
            text = ''.join(block.stripped_strings)
            if text.strip():
                lines.append(text)
        content_text = '\n'.join(lines)
    else:
        lines = []
        for p in soup.find_all('p'):
            text = ''.join(p.stripped_strings)
            if text.strip():
                lines.append(text)
        content_text = '\n\n'.join(lines)

    # 保存文本文件
    txt_path = os.path.join(save_dir, f"{base_fname}.txt")
    with open(txt_path, 'w', encoding='utf-8-sig') as f:
        f.write(content_text)

    # 保存Word文档（如果是full模式）
    docx_path = None
    if mode == 'full':
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            img_placeholder_pattern = r'(\[图片:[^\]]+\])'
            
            for para in content_text.split('\n'):
                if not para.strip():
                    continue
                    
                parts = re.split(img_placeholder_pattern, para)
                for part in parts:
                    if not part.strip():
                        continue
                    
                    img_match = re.fullmatch(r'\[图片:(.*?)\]', part.strip())
                    if img_match:
                        img_filename = img_match.group(1)
                        img_path = os.path.join(img_folder, img_filename)
                        if os.path.exists(img_path):
                            try:
                                doc.add_picture(img_path, width=Inches(4.5))
                            except:
                                doc.add_paragraph('[图片插入失败]')
                    else:
                        doc.add_paragraph(part)
            
            docx_path = os.path.join(save_dir, f"{base_fname}.docx")
            doc.save(docx_path)
            
        except ImportError:
            print("警告：未安装python-docx，跳过Word文档生成")
        except Exception as e:
            print(f"生成Word文档失败: {e}")

    return {
        "content": content_text,
        "images": images_data,
        "files": {
            "txt": f"articles/{user_info['save_dir_prefix']}/{date_str}/{base_fname}.txt",
            "docx": f"articles/{user_info['save_dir_prefix']}/{date_str}/{base_fname}.docx" if docx_path else None
        },
        "word_count": len(content_text.replace('\n', '').replace(' ', '')),
        "image_count": len(images_data)
    }

def crawl_jiuyan_article(user_key, date_str=None, try_second_time=True):
    if user_key not in JIUYAN_USERS:
        print(f"未找到用户配置: {user_key}")
        return None
    
    user_info = JIUYAN_USERS[user_key]
    
    if date_str:
        try:
            target_date = datetime.strptime(date_str, '%Y-%m-%d')
        except:
            print("日期格式错误")
            return None
    else:
        target_date = datetime.now()
    
    date_str = target_date.strftime('%Y-%m-%d')
    
    try:
        title, article_url = get_target_article_url(user_info['user_url'], date_str)
        
        if not title and try_second_time and user_info.get('retry_time'):
            print(f"未找到今日文章，将在 {user_info['retry_time']} 重试")
            return None
        
        if not title:
            print(f"未找到 {date_str} 的文章")
            return None
        
        # 获取文章内容
        soup, article_url, html = fetch_article_content(article_url)
        if not soup:
            return None
        
        # 保存文章
        save_dir = os.path.join('articles', user_info['save_dir_prefix'], date_str)
        base_fname = f"{user_info['user_name']}_{date_str}"
        
        article_data = save_article_and_generate_json(
            soup, article_url, save_dir, base_fname, user_info, date_str
        )
        
        if article_data:
            # 生成文章元数据
            metadata = {
                "title": title,
                "author": user_info['user_name'],
                "date": date_str,
                "url": article_url,
                "crawl_time": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "content_preview": article_data['content'][:200] + "..." if len(article_data['content']) > 200 else article_data['content'],
                "stats": {
                    "word_count": article_data['word_count'],
                    "image_count": article_data['image_count']
                },
                "files": article_data['files'],
                "images": article_data['images']
            }
            
            # 保存元数据
            json_path = os.path.join(save_dir, f"{base_fname}.json")
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, ensure_ascii=False, indent=2)
            
            print(f"文章已保存: {title}")
            return metadata
            
    except Exception as e:
        print(f"爬取文章时发生错误: {e}")
        return None

def update_jiuyan_index():
    """更新韭研公社文章索引"""
    index_data = {
        "update_time": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        "users": {},
        "recent_articles": []
    }
    
    articles_dir = 'articles'
    if not os.path.exists(articles_dir):
        return
    
    all_articles = []
    
    # 遍历所有用户目录
    for user_key, user_info in JIUYAN_USERS.items():
        user_dir = os.path.join(articles_dir, user_info['save_dir_prefix'])
        if not os.path.exists(user_dir):
            continue
        
        user_articles = []
        
        # 遍历日期目录
        for date_dir in os.listdir(user_dir):
            date_path = os.path.join(user_dir, date_dir)
            if not os.path.isdir(date_path):
                continue
            
            # 查找JSON文件
            json_files = [f for f in os.listdir(date_path) if f.endswith('.json')]
            for json_file in json_files:
                json_path = os.path.join(date_path, json_file)
                try:
                    with open(json_path, 'r', encoding='utf-8') as f:
                        article_data = json.load(f)
                        user_articles.append(article_data)
                        all_articles.append(article_data)
                except:
                    continue
        
        # 按日期排序
        user_articles.sort(key=lambda x: x['date'], reverse=True)
        index_data['users'][user_key] = {
            "user_name": user_info['user_name'],
            "article_count": len(user_articles),
            "latest_article": user_articles[0] if user_articles else None,
            "articles": user_articles[:10]  # 最近10篇
        }
    
    # 所有文章按时间排序
    all_articles.sort(key=lambda x: x['crawl_time'], reverse=True)
    index_data['recent_articles'] = all_articles[:20]  # 最近20篇
    
    # 保存索引
    index_path = os.path.join(articles_dir, 'index.json')
    with open(index_path, 'w', encoding='utf-8') as f:
        json.dump(index_data, f, ensure_ascii=False, indent=2)
    
    print("韭研公社文章索引已更新")

# ========== HTML生成相关函数 ==========

def generate_html():
    """生成展示页面"""
    # 读取涨停池数据索引
    limit_up_dates = []
    if os.path.exists('data/index.json'):
        with open('data/index.json', 'r', encoding='utf-8') as f:
            limit_up_dates = json.load(f)
    
    # 读取韭研公社文章索引
    jiuyan_data = {}
    jiuyan_index_path = 'articles/index.json'
    if os.path.exists(jiuyan_index_path):
        with open(jiuyan_index_path, 'r', encoding='utf-8') as f:
            jiuyan_data = json.load(f)
    
    # 生成HTML
    html_content = '''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>股市数据中心</title>
    <link rel="stylesheet" href="style.css">
</head>
<body>
    <div class="container">
        <h1>股市数据中心</h1>
        
        <div class="tabs">
            <button class="tab-button active" onclick="showTab('limitup')">涨停池数据</button>
            <button class="tab-button" onclick="showTab('jiuyan')">韭研公社文章</button>
        </div>
        
        <div id="limitup-content" class="tab-content active">
            <h2>涨停池数据</h2>
            <div class="date-list">
    '''
    
    # 添加涨停池日期列表
    for date in limit_up_dates[:30]:  # 最近30天
        html_content += f'<button class="date-button" onclick="loadLimitUpData(\'{date}\')">{date}</button>\n'
    
    html_content += '''
            </div>
            <div id="limitup-data"></div>
        </div>
        
        <div id="jiuyan-content" class="tab-content">
            <h2>韭研公社文章</h2>
    '''
    
    # 添加韭研公社内容
    if jiuyan_data:
        html_content += '<div class="jiuyan-users">'
        for user_key, user_data in jiuyan_data.get('users', {}).items():
            html_content += f'''
                <div class="user-section">
                    <h3>{user_data['user_name']} (共{user_data['article_count']}篇)</h3>
                    <div class="article-list">
            '''
            
            for article in user_data.get('articles', [])[:10]:
                html_content += f'''
                    <div class="article-item" onclick="showArticle('{article['files']['txt']}')">
                        <div class="article-title">{article['title']}</div>
                        <div class="article-meta">
                            {article['date']} | {article['stats']['word_count']}字 | {article['stats']['image_count']}图
                        </div>
                        <div class="article-preview">{article['content_preview']}</div>
                    </div>
                '''
            
            html_content += '''
                    </div>
                </div>
            '''
        html_content += '</div>'
    
    html_content += '''
        </div>
    </div>
    
    <div id="article-modal" class="modal">
        <div class="modal-content">
            <span class="close" onclick="closeModal()">&times;</span>
            <div id="article-content"></div>
        </div>
    </div>
    
    <script src="script.js"></script>
</body>
</html>
    '''
    
    with open('index.html', 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    print("HTML页面已生成")

def generate_css():
    """生成CSS样式"""
    css_content = '''
body {
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
    margin: 0;
    padding: 0;
    background-color: #f5f5f5;
}

.container {
    max-width: 1200px;
    margin: 0 auto;
    padding: 20px;
}

h1 {
    text-align: center;
    color: #333;
    margin-bottom: 30px;
}

.tabs {
    display: flex;
    justify-content: center;
    margin-bottom: 30px;
    border-bottom: 2px solid #ddd;
}

.tab-button {
    background: none;
    border: none;
    padding: 10px 30px;
    font-size: 16px;
    cursor: pointer;
    color: #666;
    transition: all 0.3s;
}

.tab-button:hover {
    color: #333;
}

.tab-button.active {
    color: #e74c3c;
    border-bottom: 2px solid #e74c3c;
}

.tab-content {
    display: none;
}

.tab-content.active {
    display: block;
}

.date-list {
    display: flex;
    flex-wrap: wrap;
    gap: 10px;
    margin-bottom: 20px;
}

.date-button {
    padding: 8px 16px;
    background: white;
    border: 1px solid #ddd;
    border-radius: 4px;
    cursor: pointer;
    transition: all 0.3s;
}

.date-button:hover {
    background: #e74c3c;
    color: white;
    border-color: #e74c3c;
}

#limitup-data {
    background: white;
    padding: 20px;
    border-radius: 8px;
    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
}

.stock-table {
    width: 100%;
    border-collapse: collapse;
}

.stock-table th,
.stock-table td {
    padding: 12px;
    text-align: left;
    border-bottom: 1px solid #eee;
}

.stock-table th {
    background: #f8f8f8;
    font-weight: 600;
}

.stock-table tr:hover {
    background: #f5f5f5;
}

.change-positive {
    color: #e74c3c;
    font-weight: 600;
}

.user-section {
    margin-bottom: 40px;
}

.user-section h3 {
    color: #333;
    margin-bottom: 15px;
}

.article-list {
    display: grid;
    gap: 15px;
}

.article-item {
    background: white;
    padding: 15px;
    border-radius: 8px;
    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    cursor: pointer;
    transition: all 0.3s;
}

.article-item:hover {
    box-shadow: 0 4px 8px rgba(0,0,0,0.15);
}

.article-title {
    font-size: 16px;
    font-weight: 600;
    color: #333;
    margin-bottom: 8px;
}

.article-meta {
    font-size: 14px;
    color: #666;
    margin-bottom: 8px;
}

.article-preview {
    font-size: 14px;
    color: #888;
    line-height: 1.5;
}

.modal {
    display: none;
    position: fixed;
    z-index: 1000;
    left: 0;
    top: 0;
    width: 100%;
    height: 100%;
    background-color: rgba(0,0,0,0.5);
}

.modal-content {
    background-color: white;
    margin: 5% auto;
    padding: 20px;
    width: 80%;
    max-width: 800px;
    max-height: 80vh;
    overflow-y: auto;
    border-radius: 8px;
}

.close {
    color: #aaa;
    float: right;
    font-size: 28px;
    font-weight: bold;
    cursor: pointer;
}

.close:hover {
    color: #000;
}

#article-content {
    line-height: 1.8;
    color: #333;
}

#article-content img {
    max-width: 100%;
    height: auto;
    margin: 10px 0;
}
    '''
    
    with open('style.css', 'w', encoding='utf-8') as f:
        f.write(css_content)
    
    print("CSS样式已生成")

def generate_js():
    """生成JavaScript"""
    js_content = '''
function showTab(tabName) {
    // 切换标签按钮状态
    const buttons = document.querySelectorAll('.tab-button');
    buttons.forEach(btn => btn.classList.remove('active'));
    event.target.classList.add('active');
    
    // 切换内容显示
    const contents = document.querySelectorAll('.tab-content');
    contents.forEach(content => content.classList.remove('active'));
    document.getElementById(tabName + '-content').classList.add('active');
}

function loadLimitUpData(date) {
    fetch(`data/${date}.json`)
        .then(response => response.json())
        .then(data => {
            displayLimitUpData(data);
        })
        .catch(error => {
            document.getElementById('limitup-data').innerHTML = '<p>加载数据失败</p>';
        });
}

function displayLimitUpData(data) {
    let html = `
        <h3>${data.date} 涨停池数据</h3>
        <p>更新时间: ${data.update_time} | 涨停数量: ${data.count}</p>
        <table class="stock-table">
            <thead>
                <tr>
                    <th>代码</th>
                    <th>名称</th>
                    <th>涨幅</th>
                    <th>价格</th>
                    <th>涨停时间</th>
                    <th>涨停原因</th>
                    <th>所属板块</th>
                </tr>
            </thead>
            <tbody>
    `;
    
    data.stocks.forEach(stock => {
        html += `
            <tr>
                <td>${stock.code}</td>
                <td>${stock.name}</td>
                <td class="change-positive">${stock.change_percent}</td>
                <td>${stock.price}</td>
                <td>${stock.limit_up_time}</td>
                <td>${stock.reason || '-'}</td>
                <td>${stock.plates || '-'}</td>
            </tr>
        `;
    });
    
    html += '</tbody></table>';
    document.getElementById('limitup-data').innerHTML = html;
}

function showArticle(filePath) {
    fetch(filePath)
        .then(response => response.text())
        .then(text => {
            document.getElementById('article-content').innerHTML = '<pre>' + text + '</pre>';
            document.getElementById('article-modal').style.display = 'block';
        })
        .catch(error => {
            alert('加载文章失败');
        });
}

function closeModal() {
    document.getElementById('article-modal').style.display = 'none';
}

// 页面加载时自动加载最新数据
window.onload = function() {
    // 尝试加载最新的涨停池数据
    fetch('data/index.json')
        .then(response => response.json())
        .then(dates => {
            if (dates.length > 0) {
                loadLimitUpData(dates[0]);
            }
        })
        .catch(error => console.error('加载索引失败'));
};
    '''
    
    with open('script.js', 'w', encoding='utf-8') as f:
        f.write(js_content)
    
    print("JavaScript已生成")

# ========== 主程序 ==========

def main():
    if len(sys.argv) < 2:
        print("使用方法:")
        print("  python scraper_fixed.py limitup  - 获取涨停池数据")
        print("  python scraper_fixed.py jiuyan   - 爬取韭研公社文章")
        print("  python scraper_fixed.py all      - 执行所有任务")
        print("  python scraper_fixed.py html     - 生成展示页面")
        print("  python scraper_fixed.py test     - 测试韭研公社连接")
        return
    
    command = sys.argv[1]
    
    if command == 'limitup':
        print("开始获取涨停池数据...")
        data = fetch_limit_up_data()
        if data:
            processed_data = process_limit_up_data(data)
            save_limit_up_data(processed_data)
    
    elif command == 'jiuyan':
        print("开始爬取韭研公社文章...")
        for user_key in JIUYAN_USERS:
            print(f"\n处理用户: {user_key}")
            crawl_jiuyan_article(user_key)
        update_jiuyan_index()
    
    elif command == 'all':
        print("执行所有任务...")
        
        # 获取涨停池数据
        print("\n1. 获取涨停池数据...")
        data = fetch_limit_up_data()
        if data:
            processed_data = process_limit_up_data(data)
            save_limit_up_data(processed_data)
        
        # 爬取韭研公社文章
        print("\n2. 爬取韭研公社文章...")
        for user_key in JIUYAN_USERS:
            print(f"\n处理用户: {user_key}")
            crawl_jiuyan_article(user_key)
        update_jiuyan_index()
        
        # 生成HTML
        print("\n3. 生成展示页面...")
        generate_html()
        generate_css()
        generate_js()
    
    elif command == 'html':
        print("生成展示页面...")
        generate_html()
        generate_css()
        generate_js()
    
    elif command == 'test':
        print("测试韭研公社连接...")
        test_url = JIUYAN_USERS['盘前纪要']['user_url']
        print(f"测试URL: {test_url}")
        
        # 使用修复后的函数测试
        title, article_url = get_target_article_url(test_url, datetime.now().strftime('%Y-%m-%d'))
        if title:
            print(f"测试成功！找到文章: {title}")
            print(f"文章URL: {article_url}")
        else:
            print("测试完成，未找到今日文章（可能今日还未发布）")
    
    else:
        print(f"未知命令: {command}")

if __name__ == "__main__":
    main()
